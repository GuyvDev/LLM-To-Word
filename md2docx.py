#!/usr/bin/env python3
"""
md2docx.py
==========
Convert a plain-text file containing Markdown headings, Hebrew (RTL) text,
and LaTeX math formulas into a natively-formatted Microsoft Word (.docx).

Architecture
------------
  python-docx   – Document container, standard runs/paragraphs/headings
  lxml          – Low-level XML element creation and tree manipulation
  docx.oxml     – python-docx helpers for namespace-aware element ops
  latex2mathml  – LaTeX string  →  W3C MathML XML string
  _convert_node – Custom recursive function: MathML element → OMML elements

Why bypass python-docx for math and RTL?
  python-docx exposes no API for:
    * <w:bidi> / <w:rtl> (paragraph/run BiDi direction)
    * <m:oMath>          (Office Math Markup Language equations)
  Both must be injected directly into the underlying lxml element tree.

Usage
-----
  python md2docx.py sample_input.txt -o result.docx   # process a file
  python md2docx.py                                    # run built-in demo
"""
from __future__ import annotations

import argparse
import io
import re
import sys
from typing import Optional
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import latex2mathml.converter


# ── XML namespace constants ────────────────────────────────────────────────────
OMML_NS   = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS      = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MATHML_NS = "http://www.w3.org/1998/Math/MathML"


def _omml(tag: str) -> str:
    """Clark-notation OMML tag, e.g. 'f' → '{…/math}f'."""
    return f"{{{OMML_NS}}}{tag}"


# ══════════════════════════════════════════════════════════════════════════════
#  PART 0a – XML property helpers
#  ────────────────────────────────
#  A paragraph's <w:pPr> and a run's <w:rPr> follow a strict schema-defined
#  child order.  Naive pPr.append() can create *duplicate* elements if the
#  active style already injected the same tag (e.g. a style-based <w:jc>).
#  Word silently ignores or misreads duplicates, which breaks RTL rendering.
#
#  _pPr_set / _rPr_set always remove the old element before inserting the new
#  one, guaranteeing there is exactly one instance of each property.
# ══════════════════════════════════════════════════════════════════════════════

def _pPr_set(pPr, tag: str, attrs: dict = None):
    """
    Insert (or replace) a <w:TAG> element inside *pPr*.

    1. Remove any existing element with the same tag to prevent duplicates.
    2. Create a fresh element, apply *attrs* (dict of qn-key → value).
    3. Append to pPr.  (Word tolerates minor schema-order deviations; the
       important correctness guarantee is uniqueness, not strict position.)
    """
    existing = pPr.find(qn(tag))
    if existing is not None:
        pPr.remove(existing)
    elem = OxmlElement(tag)
    if attrs:
        for k, v in attrs.items():
            elem.set(qn(k), v)
    pPr.append(elem)
    return elem


def _rPr_set(rPr, tag: str, attrs: dict = None):
    """Insert or replace a <w:TAG> element inside *rPr*."""
    existing = rPr.find(qn(tag))
    if existing is not None:
        rPr.remove(existing)
    elem = OxmlElement(tag)
    if attrs:
        for k, v in attrs.items():
            elem.set(qn(k), v)
    rPr.append(elem)
    return elem


# ══════════════════════════════════════════════════════════════════════════════
#  PART 0b – Script-direction text splitter
#  ─────────────────────────────────────────
#  The root cause of broken RTL in mixed-language lines is putting Hebrew and
#  Latin text in the same <w:r> run with <w:rtl/> applied globally.  That
#  forces the Latin part ("main()", formula labels, …) to be treated as RTL
#  and positioned on the wrong side of the paragraph.
#
#  _split_rtl_ltr() segments a string into alternating (text, is_rtl) pairs:
#
#    "התוכנה מתחילה עם main()"
#     ──────────────────────── ──────
#     Hebrew segment (RTL)     LTR segment
#
#  Neutral characters (spaces, digits, ASCII punctuation) inherit the direction
#  of the *current active segment*, i.e. they stay glued to whichever script
#  precedes them.  This means the space after "עם" stays in the Hebrew run,
#  which is what Unicode BiDi expects.
# ══════════════════════════════════════════════════════════════════════════════

_STRONG_RTL = re.compile(r"[\u0590-\u05FF\u0600-\u06FF\u200F]")  # Hebrew/Arabic + RLM
_STRONG_LTR = re.compile(r"[A-Za-z\u200E]")                       # Latin + LRM


def _split_rtl_ltr(text: str) -> list:
    """
    Split *text* into a list of ``(segment, is_rtl)`` tuples.

    Rules
    -----
    * Hebrew / Arabic codepoints → RTL (is_rtl=True)
    * Latin letters               → LTR (is_rtl=False)
    * Neutral chars (space, digits, punctuation) → inherit current direction.
      At the very start of the string, leading neutrals default to LTR.

    Example
    -------
    >>> _split_rtl_ltr("התוכנה מתחילה עם main()")
    [("התוכנה מתחילה עם ", True), ("main()", False)]
    """
    if not text:
        return []

    segments = []
    buf: list[str] = []
    current_rtl: bool | None = None          # None = not yet determined

    for char in text:
        if _STRONG_RTL.match(char):
            char_dir: bool | None = True
        elif _STRONG_LTR.match(char):
            char_dir = False
        else:
            char_dir = current_rtl           # neutral → inherit

        # First strong character fixes the initial direction
        if current_rtl is None and char_dir is not None:
            current_rtl = char_dir

        # Still unresolved (leading neutrals before any strong char)
        if char_dir is None or char_dir == current_rtl:
            buf.append(char)
        else:
            # Direction boundary → flush buffer, start new segment
            if buf:
                segments.append(("".join(buf), bool(current_rtl)))
            buf = [char]
            current_rtl = char_dir

    if buf:
        segments.append(("".join(buf), bool(current_rtl) if current_rtl is not None else False))

    return segments


# ══════════════════════════════════════════════════════════════════════════════
#  PART 0c – Academic document style
#  ────────────────────────────────────
#  Applies document-wide defaults once, before any content is added:
#    • Times New Roman 12pt body (or caller-specified font / size)
#    • 1.5× line spacing via WD_LINE_SPACING.ONE_POINT_FIVE
#    • Standard academic page margins (2.54 cm all sides, 3.17 cm left)
#    • Heading 1/2/3: black, proportional sizes, same font, 1.5× spacing
# ══════════════════════════════════════════════════════════════════════════════

def setup_academic_style(doc: Document,
                         base_font: str = "Times New Roman",
                         body_pt: int = 12) -> None:
    """
    Apply academic document defaults to *doc*.

    Call this immediately after ``Document()`` and before adding any content,
    so all subsequently created paragraphs inherit the correct style.
    """
    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.17)
        section.right_margin  = Cm(2.54)

    # ── Body / Normal style ──────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name         = base_font
    normal.font.size         = Pt(body_pt)
    normal.font.color.rgb    = RGBColor(0x00, 0x00, 0x00)
    pf = normal.paragraph_format
    pf.line_spacing_rule     = WD_LINE_SPACING.ONE_POINT_FIVE
    # space_after=0 so a single Enter gives exactly 1.5-line spacing — no
    # extra padding that would create large gaps between paragraphs.
    pf.space_after           = Pt(0)
    pf.space_before          = Pt(0)

    # ── Heading styles ───────────────────────────────────────────────────────
    _heading_specs = [
        #  style name    size   bold   italic  space_before
        ("Heading 1",   Pt(22), True,  False,  Pt(18)),
        ("Heading 2",   Pt(14), True,  False,  Pt(14)),
        ("Heading 3",   Pt(12), True,  True,   Pt(12)),
    ]
    for style_name, size, bold, italic, sp_before in _heading_specs:
        try:
            h = doc.styles[style_name]
        except KeyError:
            continue
        h.font.name          = base_font
        h.font.size          = size
        h.font.bold          = bold
        h.font.italic        = italic
        h.font.color.rgb     = RGBColor(0x00, 0x00, 0x00)
        hpf = h.paragraph_format
        hpf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        hpf.space_before      = sp_before
        hpf.space_after       = Pt(4)

    # ── Heading 1 special formatting: centred, underlined ────────────────────
    try:
        h1 = doc.styles["Heading 1"]
        h1.font.underline           = True
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except KeyError:
        pass


# ══════════════════════════════════════════════════════════════════════════════
#  PART 1 – RTL / BiDi paragraph helpers
#  ─────────────────────────────────────
#  Word needs two coordinated XML properties to render RTL content correctly:
#
#    Paragraph level  <w:pPr>
#      <w:bidi/>               The paragraph's *base* direction is RTL.
#                              Without this, Word ignores <w:rtl/> on runs.
#      <w:jc w:val="right"/>   Right-aligns the paragraph visually.
#
#    Run level  <w:rPr>
#      <w:rtl/>                Marks THIS run as right-to-left.
#      <w:rFonts w:cs="…">     'cs' = complex-script font slot.  Word looks
#                              here (NOT w:ascii) for Hebrew/Arabic glyphs.
#      <w:lang w:bidi="he-IL"> Spell-checker language hint.
#
#  CRITICAL: use _pPr_set / _rPr_set (not raw append) to prevent duplicates.
# ══════════════════════════════════════════════════════════════════════════════

def add_rtl_paragraph(doc: Document, text: str,
                      font_name: str = "Arial") -> None:
    """
    Add a paragraph that is purely right-to-left (e.g. a Hebrew sentence).

    The full text is placed in a single RTL run; use _add_mixed_paragraph
    for lines that mix Hebrew with Latin text.
    """
    para   = doc.add_paragraph()
    p_elem = para._element

    # ── Paragraph Properties ─────────────────────────────────────────────────
    pPr = p_elem.get_or_add_pPr()
    _pPr_set(pPr, "w:bidi", {"w:val": "1"})
    # OOXML spec §17.3.1.17: <w:bidi/> REVERSES "left"/"right" jc values.
    # "right" + bidi → physical LEFT.  "left" + bidi → physical RIGHT.
    # "start" is direction-aware and explicitly NOT reversed by bidi;
    # for RTL text it equals physical-right alignment.
    _pPr_set(pPr, "w:jc",   {"w:val": "start"})

    # ── Run ──────────────────────────────────────────────────────────────────
    run = para.add_run(text)
    _set_rtl_run(run, font_name)


def _set_rtl_run(run, font_name: str) -> None:
    """
    Inject RTL properties into *run*'s <w:rPr>.

    Called for every run that contains Hebrew/Arabic characters.

    XML injected:
        <w:rFonts w:cs="…"     ← Hebrew glyphs come from cs= slot
                  w:ascii="…"
                  w:hAnsi="…"/>
        <w:rtl/>               ← marks this run as RTL
        <w:lang w:bidi="he-IL"/>
    """
    rPr = run._r.get_or_add_rPr()
    _rPr_set(rPr, "w:rFonts", {
        "w:cs":    font_name,   # complex-script: Hebrew / Arabic glyphs
        "w:ascii": font_name,   # ASCII range fallback
        "w:hAnsi": font_name,   # high-ANSI range fallback
    })
    _rPr_set(rPr, "w:rtl",  {"w:val": "1"})
    _rPr_set(rPr, "w:lang", {"w:bidi": "he-IL"})


def _set_rtl_para(para, font_name: str = "Arial") -> None:
    """
    Mark *para* as a BiDi / RTL paragraph.

    Sets <w:bidi/> and <w:jc w:val="start"/> in <w:pPr>.
    "start" is direction-aware (= physical right for RTL) and is NOT
    reversed by <w:bidi/>, unlike the "right" literal value.
    """
    pPr = para._element.get_or_add_pPr()
    _pPr_set(pPr, "w:bidi", {"w:val": "1"})
    _pPr_set(pPr, "w:jc",   {"w:val": "start"})


# ── kept as public aliases so external callers don't break ────────────────────
def _apply_rtl_run_props(run, font_name: str = "Arial") -> None:
    _set_rtl_run(run, font_name)

def _apply_rtl_para_props(para, font_name: str = "Arial") -> None:
    _set_rtl_para(para, font_name)


# ══════════════════════════════════════════════════════════════════════════════
#  PART 2 – MathML → OMML converter
#  ──────────────────────────────────
#  Microsoft Word uses OMML (Office Math Markup Language) for equations.
#  W3C MathML (produced by latex2mathml) cannot be embedded directly; it must
#  be translated element-by-element into OMML.
#
#  Key OMML concepts used here:
#
#    <m:oMath>           Root container of an inline/display equation.
#    <m:r><m:t>…</m:t>   A math "run" – the leaf text node of an equation.
#    <m:f>               Fraction.  Children: <m:num>, <m:den>.
#    <m:rad>             Radical (root).  Children: <m:deg>, <m:e>.
#                        Set <m:degHide m:val="1"/> in <m:radPr> for √.
#    <m:sSup>            Superscript.     Children: <m:e>, <m:sup>.
#    <m:sSub>            Subscript.       Children: <m:e>, <m:sub>.
#    <m:sSubSup>         Sub+superscript. Children: <m:e>, <m:sub>, <m:sup>.
#    <m:d>               Delimiter (brackets).
#    <m:m>               Matrix/table.    Children: <m:mr>/<m:e>.
#
#  MathML → OMML mapping (per element local name):
#    mi / mn / mo / mtext  →  <m:r><m:t>…</m:t></m:r>   (leaf text run)
#    mrow / mstyle / math  →  transparent grouping (process children only)
#    mfrac                 →  <m:f>
#    msqrt                 →  <m:rad> with degHide
#    mroot                 →  <m:rad>
#    msup                  →  <m:sSup>
#    msub                  →  <m:sSub>
#    msubsup               →  <m:sSubSup>
#    mover                 →  <m:sSup>  (approximation)
#    munder                →  <m:sSub>  (approximation)
#    mfenced               →  <m:d>
#    mtable / mtr / mtd   →  <m:m> / <m:mr> / <m:e>
# ══════════════════════════════════════════════════════════════════════════════

def _me(tag: str) -> etree._Element:
    """Create a bare OMML lxml element with no children."""
    return etree.Element(_omml(tag))


def _mr(text: str) -> etree._Element:
    """
    Build a minimal OMML math run:

        <m:r>
          <m:t>TEXT</m:t>
        </m:r>

    <m:r> is the leaf node of an OMML expression; <m:t> carries the glyphs.
    """
    r = _me("r")
    t = etree.SubElement(r, _omml("t"))
    t.text = text
    return r


def _convert_node(node: etree._Element) -> list:
    """
    Recursively convert a single MathML lxml element into a list of OMML
    lxml elements.

    Returns a *list* (not a single element) because some MathML containers
    like <mrow> produce multiple peer OMML elements.
    """
    tag = etree.QName(node).localname

    # ── Leaf text nodes ───────────────────────────────────────────────────────
    if tag in ("mi", "mn", "mo", "mtext"):
        text = (node.text or "").strip()
        return [_mr(text)] if text else []

    # ── Transparent grouping ─────────────────────────────────────────────────
    if tag in ("mrow", "mstyle", "math"):
        result = []
        for child in node:
            result.extend(_convert_node(child))
        return result

    # ── Fraction: <mfrac> ────────────────────────────────────────────────────
    #
    #   OMML:    <m:f>
    #              <m:fPr><m:type m:val="bar"/></m:fPr>
    #              <m:num> … </m:num>
    #              <m:den> … </m:den>
    #            </m:f>
    #
    if tag == "mfrac":
        children = list(node)
        f    = _me("f")
        fPr  = etree.SubElement(f, _omml("fPr"))
        ftyp = etree.SubElement(fPr, _omml("type"))
        ftyp.set(_omml("val"), "bar")
        num = etree.SubElement(f, _omml("num"))
        den = etree.SubElement(f, _omml("den"))
        if len(children) >= 1:
            for elem in _convert_node(children[0]):
                num.append(elem)
        if len(children) >= 2:
            for elem in _convert_node(children[1]):
                den.append(elem)
        return [f]

    # ── Square root: <msqrt> ─────────────────────────────────────────────────
    #
    #   OMML:    <m:rad>
    #              <m:radPr><m:degHide m:val="1"/></m:radPr>
    #              <m:deg/>
    #              <m:e> … </m:e>
    #            </m:rad>
    #
    if tag == "msqrt":
        rad   = _me("rad")
        radPr = etree.SubElement(rad, _omml("radPr"))
        dh    = etree.SubElement(radPr, _omml("degHide"))
        dh.set(_omml("val"), "1")
        etree.SubElement(rad, _omml("deg"))
        e = etree.SubElement(rad, _omml("e"))
        for child in node:
            for elem in _convert_node(child):
                e.append(elem)
        return [rad]

    # ── nth root: <mroot> ────────────────────────────────────────────────────
    if tag == "mroot":
        children = list(node)
        rad  = _me("rad")
        etree.SubElement(rad, _omml("radPr"))
        deg  = etree.SubElement(rad, _omml("deg"))
        e    = etree.SubElement(rad, _omml("e"))
        if len(children) >= 2:
            for elem in _convert_node(children[1]):
                deg.append(elem)
        if len(children) >= 1:
            for elem in _convert_node(children[0]):
                e.append(elem)
        return [rad]

    # ── Superscript: <msup> ──────────────────────────────────────────────────
    if tag == "msup":
        children = list(node)
        sSup = _me("sSup")
        e    = etree.SubElement(sSup, _omml("e"))
        sup  = etree.SubElement(sSup, _omml("sup"))
        if len(children) >= 1:
            for elem in _convert_node(children[0]):
                e.append(elem)
        if len(children) >= 2:
            for elem in _convert_node(children[1]):
                sup.append(elem)
        return [sSup]

    # ── Subscript: <msub> ────────────────────────────────────────────────────
    if tag == "msub":
        children = list(node)
        sSub = _me("sSub")
        e    = etree.SubElement(sSub, _omml("e"))
        sub  = etree.SubElement(sSub, _omml("sub"))
        if len(children) >= 1:
            for elem in _convert_node(children[0]):
                e.append(elem)
        if len(children) >= 2:
            for elem in _convert_node(children[1]):
                sub.append(elem)
        return [sSub]

    # ── Sub + superscript: <msubsup> ─────────────────────────────────────────
    if tag == "msubsup":
        children = list(node)
        sss = _me("sSubSup")
        e   = etree.SubElement(sss, _omml("e"))
        sub = etree.SubElement(sss, _omml("sub"))
        sup = etree.SubElement(sss, _omml("sup"))
        if len(children) >= 1:
            for elem in _convert_node(children[0]):
                e.append(elem)
        if len(children) >= 2:
            for elem in _convert_node(children[1]):
                sub.append(elem)
        if len(children) >= 3:
            for elem in _convert_node(children[2]):
                sup.append(elem)
        return [sss]

    # ── Over / under ─────────────────────────────────────────────────────────
    if tag == "mover":
        children = list(node)
        sSup = _me("sSup")
        e    = etree.SubElement(sSup, _omml("e"))
        sup  = etree.SubElement(sSup, _omml("sup"))
        if len(children) >= 1:
            for elem in _convert_node(children[0]):
                e.append(elem)
        if len(children) >= 2:
            for elem in _convert_node(children[1]):
                sup.append(elem)
        return [sSup]

    if tag == "munder":
        children = list(node)
        sSub = _me("sSub")
        e    = etree.SubElement(sSub, _omml("e"))
        sub  = etree.SubElement(sSub, _omml("sub"))
        if len(children) >= 1:
            for elem in _convert_node(children[0]):
                e.append(elem)
        if len(children) >= 2:
            for elem in _convert_node(children[1]):
                sub.append(elem)
        return [sSub]

    # ── Delimiters / brackets: <mfenced> ─────────────────────────────────────
    #
    #   OMML:    <m:d>
    #              <m:dPr><m:begChr m:val="("/>
    #                     <m:endChr m:val=")"/></m:dPr>
    #              <m:e> … </m:e>
    #            </m:d>
    #
    if tag == "mfenced":
        open_  = node.get("open",  "(")
        close_ = node.get("close", ")")
        d    = _me("d")
        dPr  = etree.SubElement(d, _omml("dPr"))
        bChr = etree.SubElement(dPr, _omml("begChr"))
        bChr.set(_omml("val"), open_)
        eChr = etree.SubElement(dPr, _omml("endChr"))
        eChr.set(_omml("val"), close_)
        e = etree.SubElement(d, _omml("e"))
        for child in node:
            for elem in _convert_node(child):
                e.append(elem)
        return [d]

    # ── Matrix / table: <mtable> ─────────────────────────────────────────────
    if tag == "mtable":
        m_el = _me("m")
        for mtr in node:
            if etree.QName(mtr).localname != "mtr":
                continue
            mr = etree.SubElement(m_el, _omml("mr"))
            for mtd in mtr:
                if etree.QName(mtd).localname != "mtd":
                    continue
                cell = etree.SubElement(mr, _omml("e"))
                for elem in _convert_node(mtd):
                    cell.append(elem)
        return [m_el]

    # ── Fallback ──────────────────────────────────────────────────────────────
    text = "".join(node.itertext()).strip()
    return [_mr(text)] if text else []


def mathml_str_to_omml(mathml_string: str) -> etree._Element:
    """
    Convert a W3C MathML XML string to an OMML <m:oMath> lxml element.

    Pipeline
    --------
    MathML string  →  lxml.etree.fromstring()  →  mml_root
    mml_root       →  _convert_node()           →  list of OMML elements
    OMML elements  →  assembled under <m:oMath>  →  returned
    """
    mml_root = etree.fromstring(mathml_string.encode("utf-8"))
    oMath = etree.Element(_omml("oMath"), nsmap={"m": OMML_NS, "w": W_NS})
    for elem in _convert_node(mml_root):
        oMath.append(elem)
    return oMath


# ══════════════════════════════════════════════════════════════════════════════
#  PART 3 – LaTeX → OMML (public API)
# ══════════════════════════════════════════════════════════════════════════════

def latex_to_omml(latex_str: str) -> etree._Element:
    """
    Convert a LaTeX math expression to an OMML <m:oMath> lxml element.

    Pipeline: LaTeX → latex2mathml → MathML XML → mathml_str_to_omml → OMML
    """
    mathml_string = latex2mathml.converter.convert(latex_str)
    return mathml_str_to_omml(mathml_string)


def insert_latex_equation(doc: Document, latex_str: str) -> None:
    """
    Append a display-style (block) equation paragraph to *doc*.

    Word interprets <m:oMath> as a block equation when it is a direct child
    of <w:p> with no sibling <w:r> runs.  We inject it via:

        para._element.append(oMath)
    """
    para  = doc.add_paragraph()
    oMath = latex_to_omml(latex_str)
    para._element.append(oMath)


def insert_inline_latex(para, latex_str: str) -> None:
    """
    Append an inline equation to an *existing* paragraph.

    <m:oMath> is a valid sibling of <w:r> inside <w:p>.
    """
    oMath = latex_to_omml(latex_str)
    para._element.append(oMath)


# ══════════════════════════════════════════════════════════════════════════════
#  PART 4 – Markdown-aware line parser
#  ────────────────────────────────────
#  _add_mixed_paragraph is the core rendering function.  It handles any
#  combination of Hebrew text, Latin text, and inline markup on a single line.
#
#  Inline markup supported: $math$, **bold**, *italic*, `code`, ~~strike~~
#
#  Processing pipeline for "התוכנה מתחילה עם main()":
#
#    1. No inline tokens → whole line is one plain-text token.
#    2. _split_rtl_ltr() → [("התוכנה מתחילה עם ", True), ("main()", False)]
#    3. Paragraph marked BiDi (has Hebrew).
#    4. Segment 0: Hebrew run with <w:rtl/> and w:cs="Arial".
#    5. Segment 1: LTR run – NO <w:rtl/> → rendered at the visual LEFT edge
#                 of the RTL paragraph by the Unicode BiDi algorithm.
# ══════════════════════════════════════════════════════════════════════════════

_RTL_RE         = re.compile(r"[\u0590-\u05FF\u0600-\u06FF]")
_INLINE_MATH_RE = re.compile(r"\$(.+?)\$")           # legacy; kept for compat

# Full inline tokenizer: $$math$$, $math$, `code`, **bold**, *italic*, ~~strike~~
_INLINE_RE = re.compile(
    r"(\$\$[^$]+\$\$"    # $$…$$ display math used inline
    r"|\$[^$]+\$"        # $…$   inline math
    r"|`[^`]+`"          # `…`   code span
    r"|\*\*[^*]+\*\*"    # **…** bold  (must precede single-star alternative)
    r"|\*[^*]+\*"        # *…*   italic
    r"|~~[^~]+~~)"       # ~~…~~ strikethrough
)

_UL_RE = re.compile(r"^[-*+] (.+)")            # unordered list item
_OL_RE = re.compile(r"^\d+\. (.+)")            # ordered list item
_HR_RE = re.compile(r"^(\-{3,}|\*{3,}|_{3,})\s*$")  # horizontal rule


def _is_rtl(text: str) -> bool:
    return bool(_RTL_RE.search(text))


def _add_formatted_text(para, text: str, font_name: str,
                        bold: bool = False, italic: bool = False,
                        strike: bool = False, mono: bool = False) -> None:
    """
    Append *text* to *para* as one or more runs, respecting BiDi direction.

    Character-formatting flags are applied to every run produced.
    Each run covers one homogeneous script-direction segment so that Hebrew
    runs receive <w:rtl/> while Latin runs do not.
    """
    for seg_text, seg_rtl in _split_rtl_ltr(text):
        if not seg_text:
            continue
        run = para.add_run(seg_text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if strike:
            run.font.strike = True
        if mono:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        if seg_rtl:
            _set_rtl_run(run, font_name)


def _process_inline_fmt(para, text: str, font_name: str,
                        bold: bool = False, italic: bool = False,
                        strike: bool = False, mono: bool = False) -> None:
    """
    Like _process_inline() but applies character formatting to plain-text parts.

    Used so that ``**bold $math$**`` correctly renders the plain-text portions
    bold while still converting the inner ``$math$`` to an OMML equation.
    Math tokens are always passed through as-is (OMML has no bold attribute).
    """
    parts = _INLINE_RE.split(text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        if idx % 2 == 0:
            _add_formatted_text(para, part, font_name,
                                bold=bold, italic=italic,
                                strike=strike, mono=mono)
        else:
            _apply_inline_token(para, part, font_name)


def _apply_inline_token(para, token: str, font_name: str) -> None:
    """Apply a single matched _INLINE_RE token to *para*."""
    if token.startswith("$$") and token.endswith("$$"):
        latex = token[2:-2].strip()
    elif token.startswith("$") and token.endswith("$"):
        latex = token[1:-1].strip()
    else:
        latex = None

    if latex is not None:
        try:
            para._element.append(latex_to_omml(latex))
        except Exception as exc:
            para.add_run(f"[MATH ERROR: {exc}]")
    elif token.startswith("`") and token.endswith("`"):
        _add_formatted_text(para, token[1:-1], font_name, mono=True)
    elif token.startswith("**") and token.endswith("**"):
        # Recurse so inner $math$ / *italic* / `code` tokens are still processed
        _process_inline_fmt(para, token[2:-2], font_name, bold=True)
    elif token.startswith("*") and token.endswith("*"):
        _process_inline_fmt(para, token[1:-1], font_name, italic=True)
    elif token.startswith("~~") and token.endswith("~~"):
        _process_inline_fmt(para, token[2:-2], font_name, strike=True)
    else:
        _add_formatted_text(para, token, font_name)


def _process_inline(para, text: str, font_name: str) -> None:
    """
    Tokenize *text* for inline markup and append formatted runs / math to *para*.

    Token types handled
    -------------------
    $$…$$  →  OMML equation (display math used inline)
    $…$    →  OMML equation (inline math)
    `…`    →  Courier New monospace run
    **…**  →  bold run(s)
    *…*    →  italic run(s)
    ~~…~~  →  strikethrough run(s)
    plain  →  BiDi-split runs via _add_formatted_text()
    """
    parts = _INLINE_RE.split(text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        if idx % 2 == 0:
            _add_formatted_text(para, part, font_name)
        else:
            _apply_inline_token(para, part, font_name)


def _add_mixed_paragraph(doc: Document, raw_line: str,
                          font_name: str = "Arial",
                          style: str | None = None) -> None:
    """
    Add a paragraph that may contain inline markup, Hebrew text, or both.

    Parameters
    ----------
    raw_line  :  Source text.  May include $math$, **bold**, *italic*,
                 `code`, ~~strike~~, and mixed RTL/LTR text.
    font_name :  Complex-script font for Hebrew / Arabic glyphs.
    style     :  Optional paragraph style ('List Bullet', 'List Number', …).
    """
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()

    if _is_rtl(raw_line):
        _set_rtl_para(para, font_name)

    _process_inline(para, raw_line, font_name)


# ══════════════════════════════════════════════════════════════════════════════
#  PART 4b – Table, horizontal rule, and blockquote helpers
# ══════════════════════════════════════════════════════════════════════════════

def _add_horizontal_rule(doc: Document) -> None:
    """
    Add a thin horizontal line between paragraphs.

    Word has no native horizontal-rule element; the standard workaround is a
    paragraph whose <w:pBdr> bottom border spans the full text width.
    """
    para = doc.add_paragraph()
    pPr  = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "6")      # thickness in half-points
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "auto")
    pBdr.append(btm)
    pPr.append(pBdr)


def _add_blockquote(doc: Document, text: str,
                    font_name: str = "Arial") -> None:
    """
    Add a blockquote paragraph: left-indented, italic, with a grey left border.
    """
    para = doc.add_paragraph()
    if _is_rtl(text):
        _set_rtl_para(para, font_name)
    para.paragraph_format.left_indent = Cm(1.0)

    pPr  = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "888888")
    pBdr.append(left)
    pPr.append(pBdr)

    # Render content italic; inline tokens (math, bold, etc.) are passed through
    parts = _INLINE_RE.split(text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        if idx % 2 == 0:
            _add_formatted_text(para, part, font_name, italic=True)
        else:
            _apply_inline_token(para, part, font_name)


def _is_table_separator(line: str) -> bool:
    """Return True for a Markdown table separator row like ``|---|:---:|---|``."""
    cells = [c.strip() for c in line.strip("|").split("|")]
    return bool(cells) and all(re.match(r"^:?-+:?$", c) for c in cells if c)


def _add_table(doc: Document, lines: list[str], font_name: str) -> None:
    """
    Parse *lines* as a Markdown pipe table and render as a native Word table.

    Markdown pipe table format::

        | Header A | Header B | Header C |
        |----------|----------|----------|
        | cell     | cell     | cell     |

    The separator row (all dashes / colons) divides the header from the body.
    Cell content may contain any inline markup handled by _process_inline().
    """
    header_row: list[str] | None = None
    data_rows: list[list[str]] = []
    header_found = False

    for line in lines:
        if _is_table_separator(line):
            header_found = True
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if not header_found:
            header_row = cells
        else:
            data_rows.append(cells)

    all_rows: list[list[str]] = []
    if header_row is not None:
        all_rows.append(header_row)
    all_rows.extend(data_rows)

    if not all_rows:
        return

    col_count = max(len(r) for r in all_rows)
    try:
        table = doc.add_table(rows=len(all_rows), cols=col_count,
                              style="Table Grid")
    except KeyError:
        table = doc.add_table(rows=len(all_rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, cells in enumerate(all_rows):
        word_row = table.rows[row_idx]
        is_header = (row_idx == 0 and header_row is not None)
        for col_idx in range(col_count):
            cell_text = cells[col_idx] if col_idx < len(cells) else ""
            cell = word_row.cells[col_idx]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Remove the default empty run added by python-docx
            p_el = para._element
            for r_el in p_el.findall(qn("w:r")):
                p_el.remove(r_el)
            if is_header:
                run = para.add_run(cell_text)
                run.bold = True
            else:
                _process_inline(para, cell_text, font_name)


def process_text(doc: Document, text: str,
                 font_name: str = "Arial",
                 base_font: str = "Times New Roman") -> None:
    """
    Parse *text* line-by-line and add formatted content to *doc*.

    Supported markup
    ----------------
    # / ## / ###        – Headings 1–3
    $$…$$               – Display (block) equation on its own line
    $…$                 – Inline math embedded in text
    **…** / *…*         – Bold / italic
    `…`                 – Inline code (Courier New)
    ~~…~~               – Strikethrough
    - / * / +  text     – Unordered (bullet) list item
    1. / 2. …  text     – Ordered (numbered) list item
    > text              – Blockquote (indented, italic, left border)
    --- / *** / ___     – Horizontal rule
    | col | col |       – Markdown pipe table (multi-line buffer)
    Hebrew/Arabic       – Auto-detected; paragraph marked BiDi / RTL
    Mixed Hebrew+LTR    – Each script segment gets its own run (correct BiDi)
    """
    setup_academic_style(doc, base_font=base_font)

    table_buf: list[str] = []

    def _flush_table() -> None:
        if table_buf:
            _add_table(doc, table_buf, font_name)
            table_buf.clear()

    for line in text.splitlines():
        line = line.rstrip()

        # ── Table row accumulation ────────────────────────────────────────────
        if line.startswith("|"):
            table_buf.append(line)
            continue

        _flush_table()

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)

        # ── Block equations ───────────────────────────────────────────────────
        elif line.startswith("$$") and line.endswith("$$") and len(line) > 4:
            latex = line[2:-2].strip()
            try:
                insert_latex_equation(doc, latex)
            except Exception as exc:
                doc.add_paragraph(f"[MATH ERROR: {exc}]")

        # ── Horizontal rule ───────────────────────────────────────────────────
        elif _HR_RE.match(line):
            _add_horizontal_rule(doc)

        # ── Blockquote ────────────────────────────────────────────────────────
        elif line.startswith("> "):
            _add_blockquote(doc, line[2:], font_name)

        # ── Unordered list ────────────────────────────────────────────────────
        elif _UL_RE.match(line):
            content = _UL_RE.match(line).group(1)
            _add_mixed_paragraph(doc, content, font_name, style="List Bullet")

        # ── Ordered list ──────────────────────────────────────────────────────
        elif _OL_RE.match(line):
            content = _OL_RE.match(line).group(1)
            _add_mixed_paragraph(doc, content, font_name, style="List Number")

        # ── Empty line ────────────────────────────────────────────────────────
        elif not line.strip():
            doc.add_paragraph()

        # ── Regular paragraph ─────────────────────────────────────────────────
        else:
            _add_mixed_paragraph(doc, line, font_name)

    _flush_table()


# ══════════════════════════════════════════════════════════════════════════════
#  PART 5 – CLI
# ══════════════════════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════════════
#  PUBLIC API
#  ──────────────────────────────────────────────
#  convert_markdown() is the single entry-point used by the FastAPI service.
#  It keeps all I/O in memory so the web layer never touches the filesystem.
# ══════════════════════════════════════════════════════════════════════════════

def convert_markdown(
    text: str,
    font: str = "Arial",
    base_font: str = "Times New Roman",
) -> bytes:
    """
    Convert *text* (Markdown with optional Hebrew and LaTeX) to a Word .docx.

    Returns the raw .docx bytes so callers can stream it directly to HTTP
    responses without creating temporary files.

    Parameters
    ----------
    text      : Markdown source string.
    font      : Complex-script / Hebrew font (w:cs slot).  Default: Arial.
    base_font : Body and heading font.  Default: Times New Roman.
    """
    doc = Document()
    setup_academic_style(doc, base_font=base_font)
    process_text(doc, text, font_name=font, base_font=base_font)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="md2docx",
        description=(
            "Convert a Markdown/Hebrew/LaTeX text file to a Word .docx.\n"
            "Omit INPUT to run the built-in feature demo."
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    p.add_argument("input", nargs="?", metavar="INPUT",
                   help="Path to the input text file.")
    p.add_argument("-o", "--output", default="result.docx", metavar="OUTPUT",
                   help="Output filename (default: result.docx).")
    p.add_argument("--font", default="Arial", metavar="FONT",
                   help="Hebrew/complex-script font (default: Arial).")
    p.add_argument("--base-font", default="Times New Roman", metavar="BASE_FONT",
                   help="Body/heading font (default: Times New Roman).")
    return p


def main() -> None:
    args = _build_parser().parse_args()
    doc  = Document()
    setup_academic_style(doc, base_font=args.base_font)

    if args.input:
        try:
            with open(args.input, encoding="utf-8") as fh:
                text = fh.read()
        except FileNotFoundError:
            print(f"Error: file not found: {args.input}", file=sys.stderr)
            sys.exit(1)
        process_text(doc, text, font_name=args.font, base_font=args.base_font)
    else:
        print("No input file given – running built-in demo…")
        _run_demo(doc, font=args.font)

    doc.save(args.output)
    print(f"Saved -> {args.output}")


# ══════════════════════════════════════════════════════════════════════════════
#  PART 6 – Built-in demo / __main__ test case
#  ──────────────────────────────────────────────
#  Exercises every feature in isolation so the XML injection can be verified
#  without an input file:
#    1. Pure Hebrew RTL
#    2. Display block equations
#    3. Mixed Hebrew + inline OMML in the same <w:p>
#    4. Mixed BiDi: Hebrew sentence ending with LTR code token ("main()")
#    5. Inline math in English text
#    6. Additional equation examples
# ══════════════════════════════════════════════════════════════════════════════

def _run_demo(doc: Document, font: str = "Arial") -> None:
    """Populate *doc* with a full feature demonstration."""

    # ── Title ─────────────────────────────────────────────────────────────────
    doc.add_heading("md2docx – Feature Demonstration", level=1)
    doc.add_paragraph(
        "Generated entirely in Python via lxml XML injection – "
        "no image rendering, no COM automation."
    )

    # ── 1. Pure Hebrew RTL ────────────────────────────────────────────────────
    #
    # XML structure injected:
    #   <w:p>
    #     <w:pPr>
    #       <w:bidi w:val="1"/>       ← paragraph base direction = RTL
    #       <w:jc w:val="right"/>     ← right-align
    #     </w:pPr>
    #     <w:r>
    #       <w:rPr>
    #         <w:rFonts w:cs="Arial" w:ascii="Arial" w:hAnsi="Arial"/>
    #         <w:rtl w:val="1"/>      ← run is RTL
    #         <w:lang w:bidi="he-IL"/>
    #       </w:rPr>
    #       <w:t>שלום עולם…</w:t>
    #     </w:r>
    #   </w:p>
    #
    doc.add_heading("1. Hebrew RTL Paragraph", level=2)
    doc.add_paragraph("The following paragraphs are purely right-to-left:")

    add_rtl_paragraph(doc, "שלום עולם! זוהי פסקה בעברית הכתובה מימין לשמאל.",
                      font_name=font)
    add_rtl_paragraph(doc, "המתמטיקה היא שפת היקום.",
                      font_name=font)

    # ── 2. Display (block) equations ─────────────────────────────────────────
    doc.add_heading("2. Display Equations (Block)", level=2)
    doc.add_paragraph("Quadratic formula:")
    insert_latex_equation(doc, r"x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}")

    doc.add_paragraph("Energy–mass equivalence:")
    insert_latex_equation(doc, r"E = mc^2")

    doc.add_paragraph("Pythagorean theorem:")
    insert_latex_equation(doc, r"a^2 + b^2 = c^2")

    doc.add_paragraph("Derivative of sine:")
    insert_latex_equation(doc, r"\frac{d}{dx}\sin(x) = \cos(x)")

    # ── 3. Mixed Hebrew + inline OMML ─────────────────────────────────────────
    #
    # A single <w:p> containing both a Hebrew <w:r> and an <m:oMath>:
    #
    #   <w:p>
    #     <w:pPr><w:bidi …/><w:jc …/></w:pPr>
    #     <w:r>  ← Hebrew run with <w:rtl/>
    #       <w:t>הנוסחה הריבועית היא: </w:t>
    #     </w:r>
    #     <m:oMath> … </m:oMath>   ← OMML equation, sibling of the run
    #   </w:p>
    #
    doc.add_heading("3. Mixed Hebrew Text + Inline Equation", level=2)

    para = doc.add_paragraph()
    _set_rtl_para(para, font)

    run_he = para.add_run("הנוסחה הריבועית היא: ")
    _set_rtl_run(run_he, font)

    para._element.append(
        latex_to_omml(r"x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}")
    )

    # ── 4. Mixed BiDi – Hebrew sentence ending with Latin code token ──────────
    #
    # This is the key BiDi test from the screenshot.
    # "התוכנה מתחילה עם main()" produces TWO runs in one RTL paragraph:
    #
    #   Run A  "התוכנה מתחילה עם "  → <w:rtl w:val="1"/>  (Hebrew glyphs)
    #   Run B  "main()"              → NO <w:rtl/>          (LTR island)
    #
    # With <w:bidi/> on the paragraph, the Unicode BiDi algorithm places
    # Run B at the VISUAL LEFT edge of the line, matching the screenshot:
    #
    #   ┌────────────────────────────────────┐
    #   │  main()      התוכנה מתחילה עם     │
    #   └────────────────────────────────────┘
    #
    doc.add_heading("4. Mixed BiDi – Hebrew + Inline Code Token", level=2)
    doc.add_paragraph(
        "Each script segment is its own <w:r>; "
        "Latin runs carry no <w:rtl/> so BiDi positions them at the left edge:"
    )

    # _add_mixed_paragraph calls _split_rtl_ltr() internally
    _add_mixed_paragraph(doc, "התוכנה מתחילה עם main()", font_name=font)
    _add_mixed_paragraph(doc, "הפונקציה קוראת ל־initialize()", font_name=font)
    _add_mixed_paragraph(doc, "המחלקה נקראת DocumentBuilder", font_name=font)

    # ── 5. Inline math inside English text ───────────────────────────────────
    doc.add_heading("5. Inline Math in English Text", level=2)
    _add_mixed_paragraph(
        doc,
        r"When $a \neq 0$, the solutions to $ax^2 + bx + c = 0$ are "
        r"$x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}$.",
        font_name=font,
    )

    # ── 6. More equations ────────────────────────────────────────────────────
    doc.add_heading("6. Additional Equation Examples", level=2)
    for latex, label in [
        (r"x_1 + x_2 = \frac{-b}{a}",                "Sum of roots"),
        (r"x_1 \cdot x_2 = \frac{c}{a}",             "Product of roots"),
        (r"\frac{1}{1 + \frac{1}{1 + \frac{1}{x}}}", "Continued fraction"),
        (r"\sqrt{x^2 + y^2}",                         "Euclidean distance"),
    ]:
        doc.add_paragraph(f"{label}:")
        try:
            insert_latex_equation(doc, latex)
        except Exception as exc:
            doc.add_paragraph(f"  [skipped: {exc}]")

    print("Demo content built.")


if __name__ == "__main__":
    main()
